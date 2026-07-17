"""
omnistats/modules/apa_report.py
────────────────────────────────
APA 7th Edition Word document report generator.
Stage 5 (CONSOLIDATE): reads output CSVs from all prior stages.

Table order follows pipeline chronology exactly:
  Table 1 — LPA Model Fit Statistics                    (Stage 1)
  Table 2 — Profile Means, SDs, Welch ANOVA             (Stage 1)
  Table 3 — Chi-Square Tests for Demographics           (Stage 1)
  Table 4 — Profile Membership Summary                  (Stage 1)
  Table 5 — Frequentist A/B Test Results                (Stage 2)
  Table 6 — Sequential Bayesian A/B Results             (Stage 2)
  Table 7 — CUPED Variance Reduction                    (Stage 3)
  Table 8 — Full Causal Suite (DiD, IV, RDD, SCM, MC,
             BMA, CausalImpact/BSTS)                     (Stage 4)
"""
import os
import sys
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from config import OUTPUT_DIR, INDICATOR_COLS, N_PROFILES


# ─── Helpers ──────────────────────────────────────────────────────────────────

def _set_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        val = kwargs.get(side, "none")
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), val)
        if val != "none":
            b.set(qn("w:sz"), "6")
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), "000000")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _clear_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def _hline(row, pos="bottom"):
    for cell in row.cells:
        _set_border(cell, **{pos: "single"})


def _cell(cell, text, bold=False, italic=False, align="left", size=11):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    para = cell.paragraphs[0]
    para.clear()
    para.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT,
                      "center": WD_ALIGN_PARAGRAPH.CENTER,
                      "right": WD_ALIGN_PARAGRAPH.RIGHT}.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    run = para.add_run(str(text))
    run.bold = bold; run.italic = italic
    run.font.name = "Times New Roman"; run.font.size = Pt(size)


def _title(doc, text: str, num: int):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"Table {num}")
    r.bold = True; r.font.name = "Times New Roman"; r.font.size = Pt(12)
    p.add_run("\n")
    r2 = p.add_run(text)
    r2.italic = True; r2.font.name = "Times New Roman"; r2.font.size = Pt(12)


def _note(doc, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True; r.font.size = Pt(10); r.font.name = "Times New Roman"


def _table(doc, headers, rows, col_widths=None):
    ncols = len(headers)
    tbl   = doc.add_table(rows=1 + len(rows), cols=ncols)
    tbl.style = "Table Grid"
    _clear_table_borders(tbl)

    hrow = tbl.rows[0]
    for j, h in enumerate(headers):
        _cell(hrow.cells[j], h, bold=True, align="left" if j == 0 else "center")
    _hline(hrow, "top"); _hline(hrow, "bottom")

    for i, row_data in enumerate(rows):
        drow = tbl.rows[i + 1]
        for j, val in enumerate(row_data):
            _cell(drow.cells[j], val, align="left" if j == 0 else "right")

    _hline(tbl.rows[-1], "bottom")

    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[ci].width = Inches(w)

    doc.add_paragraph()


def _fmt(val, dec=2):
    if pd.isna(val):
        return "—"
    try:
        return f"{float(val):.{dec}f}"
    except Exception:
        return str(val)


# ─── Main ─────────────────────────────────────────────────────────────────────

def build_report(verbose: bool = True) -> None:
    """Build and save the full APA report to outputs/apa_report.docx."""
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # Load CSVs
    def _read(name):
        p = os.path.join(OUTPUT_DIR, name)
        return pd.read_csv(p) if os.path.exists(p) else pd.DataFrame()

    fit_df    = _read("lpa_fit_stats.csv")
    prof_df   = _read("lpa_profiles.csv")
    anov_df   = _read("anova_results.csv")
    chi2_df   = _read("chi_square_results.csv")
    ab_df     = _read("ab_test_results.csv")
    caus_df   = _read("causal_results.csv")

    doc   = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"; style.font.size = Pt(12)

    h = doc.add_heading("OmniStats — APA 7th Edition Tables", level=1)
    h.runs[0].font.name = "Times New Roman"
    doc.add_paragraph(
        "Note. Tables generated automatically by OmniStats. "
        "Significance levels: * p < .05, ** p < .01, *** p < .001."
    ).italic = True
    doc.add_page_break()

    tbl_num = 1

    # ── Table 1: LPA Fit Statistics ───────────────────────────────────────────
    _title(doc, "Model Fit Statistics for Latent Profile Analysis (K = 1–6)", tbl_num); tbl_num += 1
    if not fit_df.empty:
        hdrs = ["K", "Log-Likelihood", "AIC", "BIC", "aBIC", "Entropy", "LMR-LRT p"]
        rows = []
        for _, r in fit_df.iterrows():
            lmr = _fmt(r.get("LMR_LRT_p"), 3)
            rows.append([int(r["K"]), _fmt(r["LogLikelihood"], 2), _fmt(r["AIC"], 2),
                         _fmt(r["BIC"], 2), _fmt(r["aBIC"], 2), _fmt(r["Entropy"], 3), lmr])
        _table(doc, hdrs, rows, [0.4, 1.2, 1.0, 1.0, 1.0, 0.85, 0.9])
        _note(doc, "Note. AIC = Akaike Information Criterion; BIC = Bayesian Information Criterion; "
                   "aBIC = Sample-size Adjusted BIC; Entropy = classification entropy (0–1); "
                   "LMR-LRT p = Lo-Mendell-Rubin approximate Likelihood Ratio Test p-value.")
    doc.add_page_break()

    # ── Table 2: Profile Means + ANOVA ───────────────────────────────────────
    _title(doc, f"Profile Indicator Means (SD) and Welch's ANOVA Results (K = {N_PROFILES})", tbl_num); tbl_num += 1
    if not prof_df.empty:
        profiles = sorted(prof_df["Profile"].unique())
        hdrs = ["Indicator"] + [f"Profile {p}" for p in profiles] + ["F", "df1", "df2", "p", "η²"]
        rows = []
        for col in INDICATOR_COLS:
            if col not in prof_df.columns:
                continue
            row = [col]
            for p in profiles:
                grp = prof_df[prof_df["Profile"] == p][col].dropna()
                row.append(f"{grp.mean():.2f} ({grp.std():.2f})")
            if not anov_df.empty and col in anov_df["Indicator"].values:
                ar  = anov_df[anov_df["Indicator"] == col].iloc[0]
                sig = ar.get("sig", "")
                f_str = f"{ar['F_Welch']:.2f}{sig}" if pd.notna(ar.get("F_Welch")) else "—"
                row += [f_str, _fmt(ar["df1"], 0).replace(".00", ""),
                        _fmt(ar["df2"], 1), _fmt(ar["p"], 3), _fmt(ar["eta_squared"], 3)]
            else:
                row += ["—"] * 5
            rows.append(row)
        _table(doc, hdrs, rows)
        _note(doc, "Note. Values are M (SD) for raw scores. F = Welch's F; η² = eta-squared. * p < .05.")
    doc.add_page_break()

    # ── Table 3: Chi-Square ───────────────────────────────────────────────────
    _title(doc, "Chi-Square Tests of Profile Differences on Demographic Variables", tbl_num); tbl_num += 1
    if not chi2_df.empty:
        hdrs = ["Demographic Variable", "χ²", "df", "p", "Cramér's V"]
        rows = []
        for _, r in chi2_df.iterrows():
            rows.append([r["Demographic"],
                         f"{r['chi2']:.2f}{r.get('sig', '')}",
                         str(int(r["df"])),
                         _fmt(r["p"], 3),
                         _fmt(r["Cramers_V"], 3)])
        _table(doc, hdrs, rows, [2.0, 0.9, 0.5, 0.7, 1.1])
        _note(doc, "Note. Cramér's V is the effect size. * p < .05.")
    doc.add_page_break()

    # ── Table 4: Profile Membership ───────────────────────────────────────────
    _title(doc, "Profile Membership Counts and Percentages", tbl_num); tbl_num += 1
    if not prof_df.empty:
        total = len(prof_df)
        hdrs  = ["Profile", "n", "%", "Mean Max Prob."]
        rows  = []
        for p in sorted(prof_df["Profile"].unique()):
            grp  = prof_df[prof_df["Profile"] == p]
            prob = grp["Profile_Max_Prob"].mean() if "Profile_Max_Prob" in grp else np.nan
            rows.append([f"Profile {p}", str(len(grp)),
                         f"{len(grp) / total * 100:.1f}%",
                         f"{prob:.3f}" if not np.isnan(prob) else "—"])
        rows.append(["Total", str(total), "100.0%", "—"])
        _table(doc, hdrs, rows, [1.4, 0.7, 0.7, 1.5])
        _note(doc, "Note. Mean Max Prob. = average posterior probability of the most likely class assignment.")
    doc.add_page_break()

    # ── Table 5: A/B Test Results ─────────────────────────────────────────────
    if not ab_df.empty:
        _title(doc, "A/B Test Results Summary", tbl_num); tbl_num += 1
        cols_ab = [c for c in ["test", "p_value", "significant", "diff", "cohen_d",
                                "lift_pct", "z_stat", "t_stat"] if c in ab_df.columns]
        hdrs = [c.replace("_", " ").title() for c in cols_ab]
        rows = [[str(row[c]) for c in cols_ab] for _, row in ab_df.iterrows()]
        _table(doc, hdrs, rows)
        _note(doc, "Note. Significance at alpha = .05. Cohen's d and lift % reported where applicable.")
        doc.add_page_break()

    # ── Table 6: Bayesian A/B Results (Stage 2) ───────────────────────────────
    # Placed immediately after Frequentist A/B (Table 5):
    # both are Stage 2 — frequentist first, Bayesian second.
    bayes_path = os.path.join(OUTPUT_DIR, "bayesian_ab_results.csv")
    bayes_df   = pd.read_csv(bayes_path) if os.path.exists(bayes_path) else pd.DataFrame()
    if not bayes_df.empty:
        _title(doc, "Sequential Bayesian A/B Test Results", tbl_num); tbl_num += 1
        hdrs = ["Test", "Method", "P(B > A)", "Expected Loss",
                "95% Credible Lower", "95% Credible Upper", "ESS", "Decision"]
        rows = []
        for _, row in bayes_df.iterrows():
            rows.append([
                str(row.get("test", "—")),
                str(row.get("method", "—")),
                _fmt(row.get("p_b_beats_a"),   4),
                _fmt(row.get("expected_loss"),  6),
                _fmt(row.get("ci_lower"),        4),
                _fmt(row.get("ci_upper"),        4),
                _fmt(row.get("ess"),             1),
                str(row.get("decision", "—")),
            ])
        _table(doc, hdrs, rows)
        _note(doc,
              "Note. P(B > A) = posterior probability that Treatment exceeds Control. "
              "Expected Loss = E[max(0, θ_A − θ_B) | Data] (EVSI decision criterion). "
              "ESS = Effective Sample Size of posterior draws. "
              "Method: conjugate_beta_binomial = exact Beta-Binomial conjugate update; "
              "pymc_nuts_studentt = PyMC No-U-Turn Sampler with StudentT likelihood "
              "(robust to outliers); importance_sampling = IS fallback when PyMC unavailable. "
              "Decision threshold: P(B > A) ≥ 0.95 and Expected Loss ≤ 0.01.")
        doc.add_page_break()

    # ── Table 7: CUPED Variance Reduction (Stage 3) ─────────────────────────
    # CUPED output feeds Stage 4 Causal Inference as the adjusted outcome.
    cuped_path = os.path.join(OUTPUT_DIR, "cuped_variance_reduction.csv")
    cuped_df   = pd.read_csv(cuped_path) if os.path.exists(cuped_path) else pd.DataFrame()
    if not cuped_df.empty:
        _title(doc, "CUPED Variance Reduction (Stage 3 Pre-processing)", tbl_num); tbl_num += 1
        hdrs = ["Outcome", "Covariate", "Monotone Dir",
                "\u03b8\u0302 (slope)", "Var(Y_raw)", "Var(Y_cuped)",
                "Variance Reduction %", "Backend", "N"]
        rows = []
        for _, row in cuped_df.iterrows():
            rows.append([
                str(row.get("outcome_col",   "—")),
                str(row.get("covariate_col", "—")),
                str(row.get("monotone_dir",  "—")),
                _fmt(row.get("theta_hat"),            6),
                _fmt(row.get("var_raw"),              4),
                _fmt(row.get("var_cuped"),            4),
                _fmt(row.get("variance_reduction_pct"), 2) + "%",
                str(row.get("backend", "—")),
                str(int(row["n_obs"])) if pd.notna(row.get("n_obs")) else "—",
            ])
        _table(doc, hdrs, rows)
        _note(doc,
              "Note. CUPED (Controlled-experiment Using Pre-Experiment Data) adjusts the outcome "
              "Y_cuped = Y − θ̂(X − X̅), where X = profile_prob_max (LPA Stage 1 posterior "
              "probability of profile membership — pre-experiment, not the outcome metric itself). "
              "Monotone Dir +1 = non-decreasing profile→outcome relationship. "
              "All Stage 4 causal estimators operate on Y_cuped.")
        doc.add_page_break()

    # ── Table 8: Full Causal Suite (Stage 4) ──────────────────────────────
    # Consolidates ALL Stage 4 causal estimators (panel data + time-series)
    # into one table using the shared standardised schema.
    # Sources: causal_results.csv (DiD, IV, RDD, SCM, MC, BMA)
    #          + causal_results.csv rows from CausalImpact (appended by run_causal_suite)
    SCHEMA = ["method", "estimand", "estimate", "se",
              "ci_lower", "ci_upper", "ci_type", "p_value", "n_obs"]

    if not caus_df.empty:
        # Reload to pick up any CausalImpact rows appended by run_causal_suite
        caus_path2 = os.path.join(OUTPUT_DIR, "causal_results.csv")
        caus_df = pd.read_csv(caus_path2) if os.path.exists(caus_path2) else caus_df
        caus_df = caus_df.reindex(columns=SCHEMA)

        _title(doc, "Causal Inference Results Summary (Stage 4)", tbl_num); tbl_num += 1
        hdrs = ["Method", "Estimand", "Estimate", "SE",
                "95% CI Lower", "95% CI Upper", "CI Type", "p", "N"]
        rows = []
        for _, row in caus_df.iterrows():
            rows.append([
                str(row["method"])   if pd.notna(row.get("method"))   else "—",
                str(row["estimand"]) if pd.notna(row.get("estimand")) else "—",
                _fmt(row.get("estimate"), 4),
                _fmt(row.get("se"),       4),
                _fmt(row.get("ci_lower"), 4),
                _fmt(row.get("ci_upper"), 4),
                str(row["ci_type"]) if pd.notna(row.get("ci_type")) else "—",
                _fmt(row.get("p_value"), 4),
                str(int(row["n_obs"])) if pd.notna(row.get("n_obs")) else "—",
            ])
        _table(doc, hdrs, rows)
        _note(doc,
              "Note. All estimators operate on CUPED-adjusted outcome Y_cuped (Table 7). "
              "DiD = Callaway & Sant-Anna (2021) staggered ATT(g,t) via differences; "
              "IV = linearmodels IV2SLS (HC3 SEs; Anderson-Rubin CI when KP rk-F < 10); "
              "RDD = rdrobust CCT MSE-optimal bandwidth, bias-corrected robust CI, "
              "rddensity manipulation test; "
              "SCM = Synthetic Control Method (Abadie et al.), placebo-in-space CI; "
              "MC = Matrix Completion (Athey et al.) nuclear norm, bootstrap CI; "
              "BMA = Bayesian Model Averaging, Posterior Inclusion Probabilities for HTE; "
              "CausalImpact = Bayesian Structural Time Series (BSTS) with spike-and-slab "
              "control series selection, MCMC credible interval. "
              "CI types: doubly_robust / anderson_rubin / robust_bc / placebo_in_space / "
              "bootstrap_nuclear_norm / bsts_mcmc_credible / prophet_mcmc_credible.")

    out_path = os.path.join(OUTPUT_DIR, "apa_report.docx")
    doc.save(out_path)
    if verbose:
        print(f"[Report] APA Word document saved -> {out_path} (Tables 1–8)")
        print(f"  Table order: 1–4 = LPA | 5 = Frequentist A/B | 6 = Bayesian A/B "
              f"| 7 = CUPED | 8 = Full Causal Suite")



