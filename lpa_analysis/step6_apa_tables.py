"""
step6_apa_tables.py
───────────────────
Generate APA 7th edition formatted tables as a Word document (.docx).

Tables produced:
  1. LPA Model Fit Statistics (K = 1–6)
  2. Profile Means & SDs by Indicator + ANOVA results
  3. Chi-Square Tests for Categorical Demographics
  4. Profile Membership Summary (n and %)

Formatting rules applied:
  - No vertical borders
  - Three horizontal rules: title-bar, header-data, end
  - Times New Roman 12pt body, bold headers
  - Left-aligned label columns, centre/right-aligned numeric columns
"""
import os
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from config import OUTPUT_DIR, INDICATOR_COLS, N_PROFILES


# ── Utility helpers ──────────────────────────────────────────────────────────

def _set_cell_border(cell, **kwargs):
    """Set individual cell borders (top, bottom, left, right)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        val = kwargs.get(side, "none")
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), val)
        if val != "none":
            border.set(qn("w:sz"), "6")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "000000")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _clear_table_borders(table):
    """Remove all internal and external table-level borders."""
    tbl   = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "none")
        tblBorders.append(border)
    tblPr.append(tblBorders)


def _apply_horizontal_border(row, position="bottom"):
    """Apply a thin horizontal border to every cell in a row."""
    for cell in row.cells:
        _set_cell_border(cell, **{position: "single"})


def _cell_text(cell, text, bold=False, italic=False, align="left", font_size=11):
    """Write text to a cell with APA formatting."""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    para = cell.paragraphs[0]
    para.clear()
    align_map = {
        "left":   WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right":  WD_ALIGN_PARAGRAPH.RIGHT,
    }
    para.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    run = para.add_run(str(text))
    run.bold   = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)


def _add_table_title(doc, title_text: str, table_num: int):
    """Add APA-style table title paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(2)
    r_label = p.add_run(f"Table {table_num}")
    r_label.bold = True
    r_label.font.name = "Times New Roman"
    r_label.font.size = Pt(12)
    p.add_run("\n")
    r_title = p.add_run(title_text)
    r_title.italic = True
    r_title.font.name = "Times New Roman"
    r_title.font.size = Pt(12)


def _apa_table(doc, headers: list, rows: list, col_widths: list = None) -> None:
    """
    Create an APA-formatted table.
    headers: list of strings
    rows:    list of lists (one per data row)
    """
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.style = "Table Grid"
    _clear_table_borders(table)

    # ── Header row ────────────────────────────────────────────────────────────
    hrow = table.rows[0]
    for j, h in enumerate(headers):
        align = "left" if j == 0 else "center"
        _cell_text(hrow.cells[j], h, bold=True, align=align)
    _apply_horizontal_border(hrow, position="top")
    _apply_horizontal_border(hrow, position="bottom")

    # ── Data rows ─────────────────────────────────────────────────────────────
    for i, row_data in enumerate(rows):
        drow = table.rows[i + 1]
        for j, val in enumerate(row_data):
            align = "left" if j == 0 else "right"
            _cell_text(drow.cells[j], val, align=align)

    # ── Bottom border ─────────────────────────────────────────────────────────
    _apply_horizontal_border(table.rows[-1], position="bottom")

    # ── Column widths ─────────────────────────────────────────────────────────
    if col_widths:
        for col_idx, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[col_idx].width = Inches(width)

    doc.add_paragraph()   # spacing after table


# ── Main ──────────────────────────────────────────────────────────────────────

def build_apa_tables():
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    fit_path  = os.path.join(OUTPUT_DIR, "lpa_fit_stats.csv")
    prof_path = os.path.join(OUTPUT_DIR, "lpa_profiles.csv")
    anova_path = os.path.join(OUTPUT_DIR, "anova_results.csv")
    chi2_path  = os.path.join(OUTPUT_DIR, "chi_square_results.csv")

    fit_df  = pd.read_csv(fit_path)  if os.path.exists(fit_path)  else pd.DataFrame()
    prof_df = pd.read_csv(prof_path) if os.path.exists(prof_path) else pd.DataFrame()
    anov_df = pd.read_csv(anova_path) if os.path.exists(anova_path) else pd.DataFrame()
    chi2_df  = pd.read_csv(chi2_path)  if os.path.exists(chi2_path)  else pd.DataFrame()

    doc = Document()
    # Global font style
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # ── Title page note ───────────────────────────────────────────────────────
    h = doc.add_heading("Latent Profile Analysis — APA 7th Edition Tables", level=1)
    h.runs[0].font.name = "Times New Roman"
    doc.add_paragraph(
        "Note. Tables generated automatically from analysis outputs. "
        "Significance: * p < .05, ** p < .01, *** p < .001."
    ).italic = True
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # TABLE 1 — LPA Model Fit Statistics
    # ═══════════════════════════════════════════════════════════════════════════
    _add_table_title(doc, "Model Fit Statistics for Latent Profile Analysis (K = 1–6)", 1)

    if not fit_df.empty:
        hdrs = ["K", "Log-Likelihood", "AIC", "BIC", "aBIC", "Entropy", "LMR-LRT p"]
        rows_t1 = []
        for _, r in fit_df.iterrows():
            lmr = f"{r['LMR_LRT_p']:.3f}" if pd.notna(r.get("LMR_LRT_p")) else "—"
            rows_t1.append([
                int(r["K"]),
                f"{r['LogLikelihood']:.2f}",
                f"{r['AIC']:.2f}",
                f"{r['BIC']:.2f}",
                f"{r['aBIC']:.2f}",
                f"{r['Entropy']:.3f}",
                lmr,
            ])
        _apa_table(doc, hdrs, rows_t1, col_widths=[0.4, 1.2, 1.0, 1.0, 1.0, 0.85, 0.9])

        note = doc.add_paragraph()
        r = note.add_run(
            "Note. AIC = Akaike Information Criterion; BIC = Bayesian Information Criterion; "
            "aBIC = Sample-size Adjusted BIC; Entropy = relative entropy (0–1, higher = better separation); "
            "LMR-LRT p = Lo-Mendell-Rubin approximate Likelihood Ratio Test p-value."
        )
        r.italic = True; r.font.size = Pt(10); r.font.name = "Times New Roman"

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # TABLE 2 — Profile Means, SDs, and ANOVA Results
    # ═══════════════════════════════════════════════════════════════════════════
    _add_table_title(
        doc,
        f"Profile Indicator Means (SD) and Welch's ANOVA Results (K = {N_PROFILES})",
        2,
    )

    if not prof_df.empty:
        profiles = sorted(prof_df["Profile"].unique())
        profile_hdrs = [f"Profile {p}" for p in profiles]
        hdrs = ["Indicator"] + profile_hdrs + ["F", "df1", "df2", "p", "η²"]

        rows_t2 = []
        for col in INDICATOR_COLS:
            row_data = [col]
            for p in profiles:
                grp = prof_df[prof_df["Profile"] == p][col].dropna()
                row_data.append(f"{grp.mean():.2f} ({grp.std():.2f})")
            # ANOVA stats
            if not anov_df.empty and col in anov_df["Indicator"].values:
                ar = anov_df[anov_df["Indicator"] == col].iloc[0]
                sig = ar.get("sig", "")
                def _fmt(val, decimals=2):
                    if pd.isna(val):
                        return "—"
                    try:
                        return f"{float(val):.{decimals}f}"
                    except Exception:
                        return str(val)
                row_data += [
                    f"{ar['F_Welch']:.2f}{sig}" if pd.notna(ar['F_Welch']) else "—",
                    _fmt(ar["df1"], 0).replace(".00", ""),
                    _fmt(ar["df2"], 1),
                    _fmt(ar["p"], 3),
                    _fmt(ar["eta_squared"], 3),
                ]
            else:
                row_data += ["—"] * 5
            rows_t2.append(row_data)

        _apa_table(doc, hdrs, rows_t2)
        note = doc.add_paragraph()
        r = note.add_run(
            "Note. Values are M (SD) for raw scores. F = Welch's F; η² = eta-squared. "
            "* p < .05."
        )
        r.italic = True; r.font.size = Pt(10); r.font.name = "Times New Roman"

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # TABLE 3 — Chi-Square Tests for Categorical Demographics
    # ═══════════════════════════════════════════════════════════════════════════
    _add_table_title(doc, "Chi-Square Tests of Profile Differences on Demographic Variables", 3)

    if not chi2_df.empty:
        hdrs = ["Demographic Variable", "χ²", "df", "p", "Cramér's V"]
        rows_t3 = []
        for _, r in chi2_df.iterrows():
            rows_t3.append([
                r["Demographic"],
                f"{r['chi2']:.2f}{r.get('sig', '')}",
                str(int(r["df"])),
                f"{r['p']:.3f}",
                f"{r['Cramers_V']:.3f}",
            ])
        _apa_table(doc, hdrs, rows_t3, col_widths=[2.0, 0.9, 0.5, 0.7, 1.1])
        note = doc.add_paragraph()
        r = note.add_run(
            "Note. Cramér's V is reported as the effect size. * p < .05."
        )
        r.italic = True; r.font.size = Pt(10); r.font.name = "Times New Roman"

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # TABLE 4 — Profile Membership Summary
    # ═══════════════════════════════════════════════════════════════════════════
    _add_table_title(doc, "Profile Membership Counts and Percentages", 4)

    if not prof_df.empty:
        total = len(prof_df)
        hdrs  = ["Profile", "n", "%", "Mean Max Prob."]
        rows_t4 = []
        for p in sorted(prof_df["Profile"].unique()):
            grp    = prof_df[prof_df["Profile"] == p]
            n_p    = len(grp)
            pct    = n_p / total * 100
            avg_prob = grp["Profile_Max_Prob"].mean() if "Profile_Max_Prob" in grp else np.nan
            rows_t4.append([
                f"Profile {p}",
                str(n_p),
                f"{pct:.1f}%",
                f"{avg_prob:.3f}" if not np.isnan(avg_prob) else "—",
            ])
        rows_t4.append(["Total", str(total), "100.0%", "—"])
        _apa_table(doc, hdrs, rows_t4, col_widths=[1.4, 0.7, 0.7, 1.5])
        note = doc.add_paragraph()
        r = note.add_run(
            "Note. Mean Max Prob. = average posterior probability of the most likely class assignment."
        )
        r.italic = True; r.font.size = Pt(10); r.font.name = "Times New Roman"

    # ── Save ─────────────────────────────────────────────────────────────────
    out_path = os.path.join(OUTPUT_DIR, "apa_tables.docx")
    doc.save(out_path)
    print(f"[Step 6] APA tables saved → {out_path}")


if __name__ == "__main__":
    build_apa_tables()
